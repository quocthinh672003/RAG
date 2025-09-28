"""File toolkit using direct Python libraries for optimal performance."""

from __future__ import annotations

import io
import uuid
from typing import Any, Dict, List

import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.agent.tools.file.dto import FileToolInput, FileToolOutput
from app.agent.tools.file.storage import FileStorage


class FileToolkit:
    """File toolkit using direct Python libraries for optimal performance."""

    def __init__(self, storage: FileStorage) -> None:
        self._storage = storage

    async def generate(self, payload: FileToolInput) -> FileToolOutput:
        """Generate file using direct Python libraries."""
        file_format = payload.format.lower()
        export_id = uuid.uuid4().hex[:8]
        filename = f"export_{export_id}.{file_format}"

        # Create file content based on format
        if file_format == "pdf":
            file_content = await self._create_pdf_content(payload)
        elif file_format == "excel":
            file_content = await self._create_excel_content(payload)
        elif file_format == "word":
            file_content = await self._create_word_content(payload)
        else:
            raise ValueError(
                f"Unsupported format: {file_format}. Supported: pdf, excel, word"
            )

        # Save file
        file_path = await self._storage.save_file(
            content=file_content, filename=filename, subdir=file_format
        )

        # Get file info
        file_info = self._storage.get_file_info(file_path)

        return FileToolOutput(
            file_path=str(file_path),
            filename=filename,
            mime_type=file_info["mime_type"],
            size_bytes=file_info["size_bytes"],
            format=file_format,
        )

    async def _create_pdf_content(self, payload: FileToolInput) -> bytes:
        """Create PDF content using reportlab."""
        buffer = io.BytesIO()
        canvas_obj = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        # Add title
        y_position = height - 50
        canvas_obj.setFont("Helvetica-Bold", 16)
        title = payload.title or "Generated Document"
        canvas_obj.drawString(50, y_position, title)
        y_position -= 40

        # Add content
        canvas_obj.setFont("Helvetica", 12)
        content_lines = payload.content.split("\n")
        for line in content_lines:
            if y_position < 50:
                canvas_obj.showPage()
                y_position = height - 50
            canvas_obj.drawString(50, y_position, line[:80])
            y_position -= 20

        # Add data table if provided
        if payload.data:
            canvas_obj.showPage()
            y_position = height - 50
            canvas_obj.setFont("Helvetica-Bold", 14)
            canvas_obj.drawString(50, y_position, "Data Table")
            y_position -= 30

            canvas_obj.setFont("Helvetica", 10)
            for row in payload.data[:20]:
                if y_position < 50:
                    canvas_obj.showPage()
                    y_position = height - 50
                row_text = " | ".join([f"{k}: {v}" for k, v in row.items()])
                canvas_obj.drawString(50, y_position, row_text[:100])
                y_position -= 15

        canvas_obj.save()
        buffer.seek(0)
        return buffer.getvalue()

    async def _create_excel_content(self, payload: FileToolInput) -> bytes:
        """Create Excel content using pandas."""
        if payload.data:
            dataframe = pd.DataFrame(payload.data)
        else:
            lines = payload.content.split("\n")
            data = [{"content": line} for line in lines if line.strip()]
            dataframe = pd.DataFrame(data)

        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            dataframe.to_excel(writer, index=False, sheet_name="Data")
        buffer.seek(0)
        return buffer.getvalue()

    async def _create_word_content(self, payload: FileToolInput) -> bytes:
        """Create Word content using python-docx."""
        document = Document()
        title = payload.title or "Generated Document"
        document.add_heading(title, 0)

        if payload.content:
            document.add_paragraph(payload.content)

        if payload.data:
            document.add_heading("Data Table", level=1)
            headers = list(payload.data[0].keys()) if payload.data else []
            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"

            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header)

            for row_data in payload.data:
                row_cells = table.add_row().cells
                for i, (key, value) in enumerate(row_data.items()):
                    if i < len(row_cells):
                        row_cells[i].text = str(value)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
